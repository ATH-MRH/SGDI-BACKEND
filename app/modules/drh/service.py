from datetime import date, datetime
from io import BytesIO
from pathlib import Path
import re
import subprocess
import tempfile
import unicodedata
from types import SimpleNamespace
from typing import Any, Type

from fastapi import HTTPException
from sqlalchemy import func, or_, select, text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from app.modules.auth.models import User
from app.modules.drh.models import Candidate, Contract, ContractConditionalClause, ContractTemplate, Document, Employee, GeneratedContract, Leave, Sanction
from app.modules.irongs.models import SgdiRecord
from app.core.photo_storage import normalize_photo_fields


def _postgres_error_detail(exc: SQLAlchemyError) -> str:
    raw = str(getattr(exc, "orig", None) or exc)
    raw = " ".join(raw.split())
    return raw[:300] or exc.__class__.__name__


def list_rows(db: Session, model: Type, filters: dict[str, Any] | None = None):
    stmt = select(model)
    for key, value in (filters or {}).items():
        if value not in (None, "") and hasattr(model, key):
            stmt = stmt.where(getattr(model, key) == value)
    return db.execute(stmt.order_by(model.id.desc())).scalars().all()


def get_or_404(db: Session, model: Type, row_id: int):
    row = db.get(model, row_id)
    if not row:
        raise HTTPException(status_code=404, detail="Enregistrement introuvable")
    return row


def _clean_contract_type(value: Any) -> str:
    text = str(value or "").strip()
    return "" if text.upper() == "CDI" else text


def _parse_salary(value: Any) -> float:
    try:
        cleaned = str(value or "0").replace(" ", "").replace("\xa0", "").replace(" ", "").replace(",", ".")
        return float(cleaned) if cleaned else 0.0
    except (ValueError, TypeError):
        return 0.0


def _employee_matches_text(row: Employee, query: str) -> bool:
    if not query:
        return True
    extra = row.extra if isinstance(row.extra, dict) else {}
    haystack = " ".join(
        str(value or "")
        for value in (
            row.code,
            row.first_name,
            row.last_name,
            row.phone,
            row.email,
            row.position,
            row.society,
            extra.get("matricule"),
            extra.get("fonction"),
            extra.get("poste"),
            extra.get("affectationCourante", {}).get("siteName") if isinstance(extra.get("affectationCourante"), dict) else "",
        )
    ).lower()
    return query.lower() in haystack


def list_employees_page(
    db: Session,
    *,
    society: str | None = None,
    allowed_societies: list[str] | None = None,
    mode: str | None = None,
    q: str | None = None,
    page: int = 1,
    page_size: int = 25,
) -> dict[str, Any]:
    from sqlalchemy import or_

    page = max(int(page or 1), 1)
    page_size = min(max(int(page_size or 25), 5), 100)

    stmt = select(Employee)
    if society:
        stmt = stmt.where(Employee.society == society)
    elif allowed_societies:
        stmt = stmt.where(Employee.society.in_(allowed_societies))

    selected_mode = (mode or "actifs").strip().lower()
    if selected_mode in {"actifs", "active", "actif"}:
        stmt = stmt.where(Employee.status.in_(["actif", "active"]))
    elif selected_mode in {"absents", "absence", "absent"}:
        stmt = stmt.where(Employee.status == "absent")
    elif selected_mode in {"suspension", "suspendu", "suspendus"}:
        stmt = stmt.where(Employee.status == "suspendu")
    elif selected_mode in {"sortant", "sortants"}:
        stmt = stmt.where(Employee.status.in_(["sortant", "demissionne", "licencie"]))
    elif selected_mode not in {"all", "tous", "recap"}:
        stmt = stmt.where(Employee.status == selected_mode)

    query = str(q or "").strip()
    order = Employee.last_name.asc(), Employee.first_name.asc(), Employee.id.desc()

    if not query:
        # Chemin rapide : COUNT + OFFSET/LIMIT au niveau SQL — ne charge que la page
        total = db.scalar(select(func.count()).select_from(stmt.subquery())) or 0
        pages = max((total + page_size - 1) // page_size, 1)
        page = min(page, pages)
        items = db.execute(
            stmt.order_by(*order).offset((page - 1) * page_size).limit(page_size)
        ).scalars().all()
        return {"items": items, "total": total, "page": page, "page_size": page_size, "pages": pages}

    # Chemin avec recherche : filtre SQL large puis affinage Python sur le JSON extra
    stmt = stmt.where(
        or_(
            Employee.last_name.ilike(f"%{query}%"),
            Employee.first_name.ilike(f"%{query}%"),
            Employee.code.ilike(f"%{query}%"),
            Employee.society.ilike(f"%{query}%"),
            Employee.position.ilike(f"%{query}%"),
        )
    )
    rows = db.execute(stmt.order_by(*order)).scalars().all()
    rows = [row for row in rows if _employee_matches_text(row, query)]
    total = len(rows)
    pages = max((total + page_size - 1) // page_size, 1)
    page = min(page, pages)
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": total, "page": page, "page_size": page_size, "pages": pages}


def _candidate_is_archived(row: Candidate) -> bool:
    data = row.data if isinstance(row.data, dict) else {}
    status = str(row.status or data.get("statut") or "").strip().lower()
    return status in {"archive", "archived", "archivé", "archivee", "archivée"} or bool(
        data.get("archivedAt") or data.get("motifArchive") or data.get("commentaireArchive")
    )


def _candidate_is_reserve(row: Candidate) -> bool:
    data = row.data if isinstance(row.data, dict) else {}
    status = str(row.status or data.get("statut") or "").strip().lower()
    return status == "reserve" and bool(data.get("fichePositionValidee"))


def _candidate_is_recruited(row: Candidate) -> bool:
    data = row.data if isinstance(row.data, dict) else {}
    statuses = {
        str(row.status or "").strip().lower(),
        str(data.get("statut") or "").strip().lower(),
        str(data.get("status") or "").strip().lower(),
    }
    return bool(
        statuses & {"embauche", "embauché", "embauchee", "embauchée", "recrute", "recruté", "recrutee", "recrutée", "employe", "employé", "employee"}
        or data.get("convertedEmployeeId")
        or data.get("convertedAt")
        or data.get("employeeId")
        or data.get("agentId")
    )


def _candidate_is_active(row: Candidate) -> bool:
    return not _candidate_is_archived(row) and not _candidate_is_recruited(row)


def _candidate_matches_text(row: Candidate, query: str) -> bool:
    if not query:
        return True
    data = row.data if isinstance(row.data, dict) else {}
    haystack = " ".join(
        str(value or "")
        for value in (
            row.first_name,
            row.last_name,
            row.phone,
            row.email,
            row.desired_position,
            row.society,
            data.get("nom"),
            data.get("prenom"),
            data.get("telephone"),
            data.get("posteSouhaite"),
            data.get("wilaya"),
        )
    ).lower()
    return query.lower() in haystack


def list_candidates_page(
    db: Session,
    *,
    society: str | None = None,
    allowed_societies: list[str] | None = None,
    mode: str | None = None,
    q: str | None = None,
    page: int = 1,
    page_size: int = 25,
) -> dict[str, Any]:
    stmt = select(Candidate)
    if society:
        stmt = stmt.where(Candidate.society == society)
    elif allowed_societies:
        stmt = stmt.where(Candidate.society.in_(allowed_societies))

    rows = db.execute(stmt.order_by(Candidate.id.desc())).scalars().all()
    selected_mode = (mode or "").strip().lower()
    if selected_mode in {"archive", "archived", "archives"}:
        rows = [row for row in rows if _candidate_is_archived(row)]
    elif selected_mode in {"reserve", "reserves"}:
        rows = [row for row in rows if _candidate_is_active(row) and _candidate_is_reserve(row)]
    elif selected_mode in {"new", "nouveau", "nouvelle", "recrutement"}:
        rows = [row for row in rows if _candidate_is_active(row) and not _candidate_is_reserve(row)]

    query = str(q or "").strip()
    if query:
        rows = [row for row in rows if _candidate_matches_text(row, query)]

    page = max(int(page or 1), 1)
    page_size = min(max(int(page_size or 25), 5), 100)
    total = len(rows)
    pages = max((total + page_size - 1) // page_size, 1)
    if page > pages:
        page = pages
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": total, "page": page, "page_size": page_size, "pages": pages}



def _candidate_text(value: Any) -> str:
    return str(value or "").strip()

def _candidate_data(values: dict[str, Any], existing: Candidate | None = None) -> dict[str, Any]:
    data = values.get("data")
    if isinstance(data, dict):
        return dict(data)
    if existing and isinstance(existing.data, dict):
        return dict(existing.data)
    return {}

def _candidate_status(value: Any) -> str:
    return str(value or "nouvelle").strip().lower()

def _candidate_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "oui", "yes", "on"}

def _candidate_sections(data: dict[str, Any]) -> dict[str, Any]:
    raw = data.get("sectionValidations")
    return raw if isinstance(raw, dict) else {}


def _normalized_list(values: Any) -> list[str]:
    if not isinstance(values, list):
        return []
    return [str(v).strip() for v in values if str(v).strip()]


def _user_society_allowed(user: User, society: str | None) -> bool:
    allowed = _normalized_list(user.authorized_societies)
    return not allowed or bool(society and society in allowed)


def _user_matches_module(user: User, module: str) -> bool:
    role = str(user.role or "").strip().lower()
    level = str(user.access_level or "").strip().lower()
    structures = {s.lower() for s in _normalized_list(user.authorized_structures)}
    if module in structures or "admin" in structures:
        return True
    if role in {"admin", "adm", "adm1", "adm2"} or level.startswith("adm"):
        return True
    if module == "drh":
        return role in {"rh", "drh"}
    if module == "ops":
        return role in {"ops", "dispatch"}
    if module == "materiel":
        return role in {"ops", "dispatch"}
    if module == "commercial":
        return role in {"commercial"}
    return False


def _next_sgdi_position(db: Session, collection: str) -> int:
    current = db.scalar(select(func.max(SgdiRecord.position)).where(SgdiRecord.collection == collection))
    return int(current or 0) + 1


def _upsert_sgdi_item(db: Session, collection: str, item_id: str, data: dict[str, Any]) -> None:
    row = db.execute(select(SgdiRecord).where(SgdiRecord.collection == collection, SgdiRecord.item_id == item_id)).scalar_one_or_none()
    label = str(data.get("title") or data.get("sujet") or data.get("message") or item_id)
    if row:
        row.data = data
        row.label = label
        return
    db.add(SgdiRecord(collection=collection, item_id=item_id, position=_next_sgdi_position(db, collection), kind="item", data=data, label=label))


def _create_candidate_workflow(db: Session, row: Candidate, username: str | None) -> None:
    data = row.data if isinstance(row.data, dict) else {}
    society = row.society or data.get("societe") or data.get("society")
    candidate_name = f"{row.last_name or data.get('nom') or ''} {row.first_name or data.get('prenom') or ''}".strip() or "Candidat"
    created_at = datetime.utcnow().isoformat()
    modules = [
        ("drh", "Suivi contractualisation", "Contrôler le passage du candidat vers la contractualisation.", "contrats/a_contractualiser"),
        ("ops", "Affectation OPS à préparer", "Préparer le site, le poste et l'affectation opérationnelle.", "effectif/instance_affectation"),
        ("materiel", "Dotation matériel à préparer", "Préparer la dotation initiale selon le poste demandé.", "materiel/dotation"),
        ("commercial", "Information recrutement validé", "Prendre connaissance du recrutement validé et de son impact client/site si nécessaire.", "commercial/dashboard"),
    ]
    for module, title, message, route in modules:
        task_id = f"fiche_position_{row.id}_{module}"
        task = {
            "id": task_id,
            "type": "fiche_position_validee",
            "module": module,
            "title": title,
            "message": message,
            "status": "open",
            "priority": "high" if module in {"ops", "materiel"} else "normal",
            "societe": society,
            "candidateId": row.data.get("id") if isinstance(row.data, dict) else "",
            "candidateBackendId": row.id,
            "candidateName": candidate_name,
            "poste": data.get("posteSouhaite") or data.get("posteContrat") or row.desired_position,
            "route": route,
            "createdAt": created_at,
            "createdBy": username or "system",
        }
        _upsert_sgdi_item(db, "workflowTasks", task_id, task)

    users = db.execute(select(User).where(User.is_active.is_(True))).scalars().all()
    for user in users:
        if user.username == username:
            continue
        if not _user_society_allowed(user, society):
            continue
        if not any(_user_matches_module(user, module) for module, *_ in modules):
            continue
        msg_id = f"fiche_position_{row.id}_{user.username}"
        message = (
            f"Nouvelle fiche de position validée : {candidate_name}."
            f"\nSociété : {society or 'Non renseignée'}"
            f"\nPoste : {data.get('posteSouhaite') or data.get('posteContrat') or row.desired_position or 'Non renseigné'}"
            "\nConsultez votre tableau de bord pour traiter la tâche correspondante."
        )
        _upsert_sgdi_item(db, "echanges", msg_id, {
            "id": msg_id,
            "date": created_at,
            "from": username or "system",
            "to": user.username,
            "sujet": "Fiche de position validée",
            "importance": "Élevée",
            "message": message,
            "attachments": [],
            "obligation": "Traitement selon module",
            "droit": society or "",
            "conduite": "Consulter tableau de bord",
            "receivedBy": [],
            "luPar": [username or "system"],
            "type": "message",
            "workflowTaskType": "fiche_position_validee",
            "candidateBackendId": row.id,
            "societe": society,
        })

def _record_candidate_recruitment_creation(db: Session, row: Candidate, username: str | None) -> None:
    data = row.data if isinstance(row.data, dict) else {}
    created_at = datetime.utcnow().isoformat()
    society = row.society or data.get("societe") or data.get("society")
    candidate_name = f"{row.last_name or data.get('nom') or ''} {row.first_name or data.get('prenom') or ''}".strip() or "Candidat"
    row.data = normalize_photo_fields(
        {
            **data,
            "moduleOrigine": data.get("moduleOrigine") or "recrutement",
            "recruitmentAction": "fiche_position_created",
            "fichePositionCreatedAt": data.get("fichePositionCreatedAt") or created_at,
            "fichePositionCreatedBy": data.get("fichePositionCreatedBy") or username or "system",
        },
        fallback=str(row.id),
    )
    activity_id = f"recrutement_fiche_position_{row.id}"
    _upsert_sgdi_item(db, "activityLog", activity_id, {
        "id": activity_id,
        "date": created_at,
        "user": username or "system",
        "module": "recrutement",
        "action": "Création fiche de position",
        "details": (
            f"Nouvelle fiche de position / recrutement : {candidate_name}"
            f" · Société : {society or 'Non renseignée'}"
            f" · Poste : {data.get('posteSouhaite') or row.desired_position or 'Non renseigné'}"
        ),
        "candidateBackendId": row.id,
        "candidateName": candidate_name,
        "societe": society,
        "type": "recrutement",
    })

def _candidate_required_fields(section: str) -> list[tuple[str, str]]:
    return {
        "identification": [
            ("nom", "Nom"),
            ("prenom", "Prénom"),
            ("dateNaissance", "Date de naissance"),
            ("lieuNaissance", "Lieu de naissance"),
            ("sexe", "Sexe"),
            ("nomPere", "Nom du père"),
            ("nomMere", "Nom de la mère"),
            ("nin", "NIN"),
            ("situation", "Situation familiale"),
            ("source", "Source"),
        ],
        "poste": [("posteSouhaite", "Poste souhaité"), ("telephone", "Téléphone")],
        "avis": [
            ("avisDecision", "Décision"),
            ("avisDate", "Date de l'avis"),
            ("avisRecruteur", "Recruteur"),
            ("avisCommentaire", "Commentaire"),
        ],
        "contact": [
            ("adresse", "Adresse"),
            ("commune", "Commune"),
            ("wilaya", "Wilaya"),
            ("contactUrgenceLien", "Lien contact urgence"),
            ("contactUrgenceNom", "Nom contact urgence"),
            ("contactUrgenceTel", "Téléphone urgence"),
        ],
    }.get(section, [])

def _candidate_age_on_save(date_naissance: Any) -> int | None:
    if not date_naissance:
        return None
    try:
        born = date.fromisoformat(str(date_naissance)[:10])
    except ValueError:
        raise HTTPException(status_code=422, detail="Date de naissance invalide")
    today = date.today()
    return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

def _validate_candidate_form_rules(values: dict[str, Any], existing: Candidate | None = None, section: str | None = None) -> None:
    data = _candidate_data(values, existing)
    first_name = _candidate_text(values.get("first_name", existing.first_name if existing else data.get("prenom")))
    last_name = _candidate_text(values.get("last_name", existing.last_name if existing else data.get("nom")))
    if len(first_name) < 2 or len(last_name) < 2:
        raise HTTPException(status_code=422, detail="Nom et prénom obligatoires pour créer une candidature.")

    nin = _candidate_text(data.get("nin"))
    nin_alert_authorized = _candidate_bool(data.get("importNinWarning")) or _candidate_bool(data.get("ninAlert"))
    if nin and not re.fullmatch(r"\d{10}", nin) and not nin_alert_authorized:
        raise HTTPException(status_code=422, detail="Le NIN doit contenir exactement 10 chiffres")
    age = _candidate_age_on_save(data.get("dateNaissance"))
    age_alert_authorized = _candidate_bool(data.get("importAgeWarning")) or _candidate_bool(data.get("ageAlert"))
    if age is not None and age < 20 and not age_alert_authorized:
        raise HTTPException(status_code=422, detail="Le candidat doit avoir au moins 20 ans à la date d'enregistrement")

    if section:
        missing = [label for field, label in _candidate_required_fields(section) if not _candidate_text(data.get(field))]
        if missing:
            raise HTTPException(status_code=422, detail="Champs obligatoires manquants : " + ", ".join(missing))

def _validate_candidate_transition(values: dict[str, Any], existing: Candidate | None = None) -> None:
    data = _candidate_data(values, existing)
    status_value = values.get("status", existing.status if existing else "nouvelle")
    status_norm = _candidate_status(status_value)
    sections = _candidate_sections(data)
    all_sections = ["identification", "mensurations", "militaire", "poste", "avis", "contact", "habilitations", "experience"]
    etape1 = ["identification", "mensurations", "militaire", "poste", "avis"]

    if _candidate_bool(data.get("fichePositionValidee")) or status_norm in {"reserve", "a_contractualiser", "embauche"}:
        missing = [key for key in all_sections if not sections.get(key)]
        if missing:
            raise HTTPException(status_code=422, detail="Fiche de position refusée : sections non validées (" + ", ".join(missing) + ")")

    if status_norm == "reserve":
        missing = [key for key in etape1 if not sections.get(key)]
        if missing:
            raise HTTPException(status_code=422, detail="Mise en réserve refusée : étape précédente non validée")

def _candidate_values(payload: Any, existing: Candidate | None = None, partial: bool = False) -> dict[str, Any]:
    values = payload.model_dump(exclude_unset=True)
    data = values.get("data")
    if isinstance(data, dict):
        raw_id = str(data.get("id") or "").strip()
        if raw_id.startswith("tmp_cd_"):
            raise HTTPException(status_code=422, detail="Candidature temporaire refusée. Enregistrez uniquement une fiche complète.")
        data.pop("isNew", None)
        values["data"] = normalize_photo_fields(data, fallback=raw_id or values.get("last_name") or "candidate")
    first_name = _candidate_text(values.get("first_name", existing.first_name if existing else ""))
    last_name = _candidate_text(values.get("last_name", existing.last_name if existing else ""))
    if len(first_name) < 2 or len(last_name) < 2:
        raise HTTPException(status_code=422, detail="Nom et prénom obligatoires pour créer une candidature.")
    if not partial or "first_name" in values:
        values["first_name"] = first_name
    if not partial or "last_name" in values:
        values["last_name"] = last_name
    _validate_candidate_form_rules(values, existing)
    _validate_candidate_transition(values, existing)
    return values

def validate_candidate_section(db: Session, payload: Any, section: str, existing_id: int | None = None, username: str | None = None):
    values = payload.model_dump(exclude_unset=True)
    existing = get_or_404(db, Candidate, existing_id) if existing_id else None
    data = _candidate_data(values, existing)
    order = ["identification", "mensurations", "militaire", "poste", "avis", "contact", "habilitations", "experience"]
    if section not in order:
        raise HTTPException(status_code=422, detail="Section inconnue")
    sections = dict(_candidate_sections(data))
    idx = order.index(section)
    previous_missing = [key for key in order[:idx] if not sections.get(key)]
    if previous_missing:
        raise HTTPException(status_code=422, detail="Validez d'abord la section précédente : " + ", ".join(previous_missing))
    _validate_candidate_form_rules(values, existing, section=section)
    sections[section] = {"by": username or "system", "at": datetime.utcnow().isoformat()}
    data["sectionValidations"] = sections
    values["data"] = data
    return {"status": "success", "data": {"section": section, "sectionValidations": sections, "next": order[idx + 1] if idx + 1 < len(order) else None}}

def marquer_a_contractualiser(db: Session, candidate_id: int, username: str | None = None):
    row = get_or_404(db, Candidate, candidate_id)
    if row.status not in ("reserve", "a_contractualiser"):
        raise HTTPException(status_code=422, detail="Le candidat doit être en réserve pour être contractualisé")
    data = row.data if isinstance(row.data, dict) else {}
    if not data.get("fichePositionValidee"):
        raise HTTPException(status_code=422, detail="La fiche de position doit être validée avant la contractualisation")
    row.status = "a_contractualiser"
    row.data = {
        **data,
        "statut": "a_contractualiser",
        "contractualisationAt": datetime.utcnow().isoformat(),
        "contractualisationBy": username or "system",
    }
    db.commit()
    db.refresh(row)
    return row


def validate_candidate_final(db: Session, candidate_id: int, username: str | None = None):
    row = get_or_404(db, Candidate, candidate_id)
    data = row.data if isinstance(row.data, dict) else {}
    sections = _candidate_sections(data)
    order = ["identification", "mensurations", "militaire", "poste", "avis", "contact", "habilitations", "experience"]
    missing = [key for key in order if not sections.get(key)]
    if missing:
        raise HTTPException(status_code=422, detail="Fiche candidat refusée : sections non validées (" + ", ".join(missing) + ")")
    values = {
        "first_name": row.first_name,
        "last_name": row.last_name,
        "status": "reserve",
        "data": {
            **data,
            "statut": "reserve",
            "fichePositionValidee": True,
            "fichePositionValideeAt": datetime.utcnow().isoformat(),
            "fichePositionValideeBy": username or "system",
        },
    }
    _validate_candidate_form_rules(values, row)
    _validate_candidate_transition(values, row)
    row.status = "reserve"
    row.data = normalize_photo_fields(values["data"], fallback=str(row.id))
    _create_candidate_workflow(db, row, username)
    db.commit()
    db.refresh(row)
    return row

def create_candidate(db: Session, payload: Any, username: str | None = None):
    row = Candidate(**_candidate_values(payload))
    db.add(row)
    db.flush()
    _record_candidate_recruitment_creation(db, row, username)
    db.commit()
    db.refresh(row)
    return row

def update_candidate(db: Session, candidate_id: int, payload: Any):
    row = get_or_404(db, Candidate, candidate_id)
    for key, value in _candidate_values(payload, existing=row, partial=True).items():
        setattr(row, key, value)
    db.commit()
    db.refresh(row)
    return row

def create_row(db: Session, model: Type, payload: Any):
    values = payload.model_dump(exclude_unset=True)
    if model is Employee and isinstance(values.get("extra"), dict):
        values["extra"] = normalize_photo_fields(values["extra"], fallback=values.get("code") or "employee")
    row = model(**values)
    db.add(row)
    try:
        db.commit()
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Enregistrement PostgreSQL refusé : {_postgres_error_detail(exc)}") from exc
    db.refresh(row)
    return row


def update_row(db: Session, model: Type, row_id: int, payload: Any):
    row = get_or_404(db, model, row_id)
    values = payload.model_dump(exclude_unset=True)
    if model is Employee and isinstance(values.get("extra"), dict):
        values["extra"] = normalize_photo_fields(values["extra"], fallback=getattr(row, "code", None) or str(row_id))
    for key, value in values.items():
        setattr(row, key, value)
    try:
        db.commit()
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Enregistrement PostgreSQL refusé : {_postgres_error_detail(exc)}") from exc
    db.refresh(row)
    return row


def delete_row(db: Session, model: Type, row_id: int):
    row = get_or_404(db, model, row_id)
    db.delete(row)
    db.commit()
    return {"deleted": True, "id": row_id}


def delete_employee_direct(db: Session, employee_id: int):
    employee = get_or_404(db, Employee, employee_id)
    refs = {str(employee_id), str(employee.code or "").strip()}
    extra = employee.extra if isinstance(employee.extra, dict) else {}
    legacy = extra.get("_legacy") if isinstance(extra.get("_legacy"), dict) else {}
    refs.update(str(legacy.get(key) or "").strip() for key in ("id", "backendId", "matricule", "code"))
    refs = {value for value in refs if value}
    linked_collections = {
        "agents", "employees", "conges", "contrats", "contratsPersonnel", "materiel",
        "pointages", "pointageMensuel", "feuillePresence", "demandesPersonnel",
        "demandesStructure", "missions", "siteInspections", "stockMouvements",
    }
    linked_fields = {
        "id", "backendId", "agentId", "employeeId", "employee_id",
        "beneficiaireAgentId", "retourAgentId", "matricule", "code",
    }

    def matches(row: SgdiRecord) -> bool:
        if str(row.item_id or "").strip() in refs:
            return True
        if not isinstance(row.data, dict):
            return False
        return any(str(row.data.get(field) or "").strip() in refs for field in linked_fields)

    deleted_legacy = 0
    rows = db.execute(select(SgdiRecord).where(SgdiRecord.collection.in_(linked_collections))).scalars().all()
    for row in rows:
        if matches(row):
            db.delete(row)
            deleted_legacy += 1

    db.delete(employee)
    db.commit()
    return {"deleted": True, "id": employee_id, "storage": "sql", "legacy_deleted": deleted_legacy}


def drh_dashboard(db: Session):
    return {
        "employees_total": 0,
        "employees_by_status": {},
        "candidates_by_status": {},
        "leaves_pending": 0,
        "trial_periods": [],
    }


def recruit_candidate(db: Session, candidate_id: int):
    candidate = get_or_404(db, Candidate, candidate_id)
    _validate_candidate_transition({"status": "embauche", "data": candidate.data or {}}, candidate)
    data = candidate.data if isinstance(candidate.data, dict) else {}
    next_code = next_employee_code(db, candidate.society)
    recruit_date = date.today()

    # typeContrat stocké tel quel (CDD/CDI/CIDD…) — ne pas passer par _clean_contract_type
    # qui efface CDI. posteContrat est le poste de travail, pas le type de contrat.
    raw_type = str(data.get("typeContrat") or "").strip().upper()
    contract_type = raw_type if raw_type in {"CDD", "CDI", "CIDD", "CTA", "APPRENTISSAGE"} else "CDD"

    # Période d'essai selon le type (convention sécurité privée algérienne)
    trial_days = {"CDI": 90, "CDD": 30, "CIDD": 45}.get(contract_type, 30)
    trial_end = date.fromordinal(recruit_date.toordinal() + trial_days)

    # Date de fin de contrat (CDD uniquement, depuis les données candidat si dispo)
    raw_end = str(data.get("dateFinContrat") or "").strip()
    contract_end: date | None = None
    if contract_type == "CDD":
        try:
            contract_end = date.fromisoformat(raw_end[:10]) if raw_end else None
        except ValueError:
            contract_end = None

    employee = Employee(
        code=next_code,
        first_name=candidate.first_name,
        last_name=candidate.last_name,
        phone=candidate.phone,
        email=candidate.email,
        position=candidate.desired_position,
        society=candidate.society,
        salary_net=candidate.expected_salary or 0,
        contract_type=contract_type,
        status="actif",
        recruit_date=recruit_date,
        trial_end_date=trial_end,
        contract_end_date=contract_end,
        extra=candidate.data or {},
    )
    candidate.status = "embauche"
    candidate.data = {
        **(candidate.data if isinstance(candidate.data, dict) else {}),
        "statut": "embauche",
        "status": "embauche",
        "convertedAt": datetime.utcnow().isoformat(),
    }
    db.add(employee)
    db.flush()
    contract = Contract(
        employee_id=employee.id,
        contract_type=contract_type,
        position=candidate.desired_position,
        start_date=recruit_date,
        end_date=contract_end,
        trial_end_date=trial_end,
        salary_net=candidate.expected_salary or 0,
        status="actif",
    )
    db.add(contract)
    db.commit()
    db.refresh(employee)
    return employee


def approve_leave(db: Session, leave_id: int):
    leave = get_or_404(db, Leave, leave_id)
    leave.status = "approuve"
    leave.decided_at = datetime.utcnow()
    db.commit()
    db.refresh(leave)
    return leave


def refuse_leave(db: Session, leave_id: int):
    leave = get_or_404(db, Leave, leave_id)
    leave.status = "refuse"
    leave.decided_at = datetime.utcnow()
    db.commit()
    db.refresh(leave)
    return leave


def fiche_position(db: Session, employee_id: int):
    from app.modules.materiel.models import EmployeeEquipment
    from app.modules.ops.models import Assignment, DailyPresence, Event

    employee = get_or_404(db, Employee, employee_id)
    contracts = list_rows(db, Contract, {"employee_id": employee_id})
    leaves = list_rows(db, Leave, {"employee_id": employee_id})
    sanctions = list_rows(db, Sanction, {"employee_id": employee_id})
    documents = list_rows(db, Document, {"owner_type": "employee", "owner_id": employee_id})
    assignments = list_rows(db, Assignment, {"employee_id": employee_id})
    pointage = list_rows(db, DailyPresence, {"employee_id": employee_id})
    events = list_rows(db, Event, {"employee_id": employee_id})
    equipment = list_rows(db, EmployeeEquipment, {"employee_id": employee_id, "status": "attribue"})
    return {
        "employee": employee,
        "contracts": contracts,
        "leaves": leaves,
        "sanctions": sanctions,
        "documents": documents,
        "assignments": assignments,
        "pointage": pointage,
        "events": events,
        "equipment": equipment,
    }



DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"
DEFAULT_APS_TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "static" / "contract_templates" / "contrat_aps_2026_sgdi.docx"
PLACEHOLDER_RE = re.compile(r"{{\s*([A-Z0-9_]+)\s*}}")


def extract_docx_placeholders(content: bytes) -> list[str]:
    try:
        from docx import Document as WordDocument
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="Dépendance python-docx manquante") from exc
    try:
        doc = WordDocument(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Modèle Word .docx invalide") from exc
    text_parts: list[str] = []
    text_parts.extend(p.text or "" for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.extend(p.text or "" for p in cell.paragraphs)
    found = sorted({m.group(1).upper() for m in PLACEHOLDER_RE.finditer("\n".join(text_parts))})
    return found


def _date_str(value: Any) -> str:
    if isinstance(value, (date, datetime)):
        return value.strftime("%d/%m/%Y")
    return str(value or "")


EMPLOYEE_CODE_SERIE_LIMIT = 200


def _society_code_key(value: Any) -> str:
    normalized = unicodedata.normalize("NFD", str(value or "").strip())
    without_accents = "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")
    return " ".join(without_accents.upper().split())


def employee_code_prefixes_for_society(society: Any) -> list[str]:
    key = _society_code_key(society)
    if key == "IRON GLOBAL SECURITE":
        return ["A", "B", "C"]
    if key == "IRON GLOBAL SOLUTION":
        return ["K", "W"]
    if key == "SWORD CORPORATION":
        return ["S"]
    if key == "SWORD CONSTRUCTION":
        return ["T"]
    return list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")


def _employee_code_society_priority(key: str) -> tuple[int, str]:
    known = {
        "IRON GLOBAL SECURITE": 0,
        "IRON GLOBAL SOLUTION": 1,
        "SWORD CORPORATION": 2,
        "SWORD CONSTRUCTION": 3,
    }
    return (known.get(key, 99), key)


def _employee_code_sort_key(employee: Employee) -> tuple[str, str, int]:
    return (
        str(employee.last_name or "").strip().upper(),
        str(employee.first_name or "").strip().upper(),
        int(employee.id),
    )


def _employee_code_extra(extra: Any, code: str) -> Any:
    if not isinstance(extra, dict):
        return extra
    data = dict(extra)
    data["matricule"] = code
    data["code"] = code
    legacy = data.get("_legacy")
    if isinstance(legacy, dict):
        legacy = dict(legacy)
        legacy["matricule"] = code
        legacy["code"] = code
        data["_legacy"] = legacy
    return data


def _employee_code_extra_matches(extra: Any, code: str) -> bool:
    if not isinstance(extra, dict):
        return True
    if str(extra.get("matricule") or "").strip().upper() != code:
        return False
    if str(extra.get("code") or "").strip().upper() != code:
        return False
    legacy = extra.get("_legacy")
    if isinstance(legacy, dict):
        if str(legacy.get("matricule") or "").strip().upper() != code:
            return False
        if str(legacy.get("code") or "").strip().upper() != code:
            return False
    return True


def _employee_code_sequence(prefixes: list[str], count: int, used: set[str]) -> list[str]:
    codes: list[str] = []
    for prefix in prefixes:
        for number in range(1, EMPLOYEE_CODE_SERIE_LIMIT + 1):
            code = f"{prefix}{number:02d}"
            if code in used:
                continue
            codes.append(code)
            if len(codes) == count:
                return codes
    raise HTTPException(
        status_code=409,
        detail=f"Série de codes employé saturée ({', '.join(prefixes)} 01-{EMPLOYEE_CODE_SERIE_LIMIT})",
    )


def repair_employee_codes_if_needed(db: Session) -> int:
    if db.bind.dialect.name == "postgresql":
        db.execute(text("SELECT pg_advisory_xact_lock(20260612)"))
    employees = db.execute(select(Employee).order_by(Employee.id.asc())).scalars().all()
    grouped: dict[str, tuple[str, list[Employee]]] = {}
    for employee in employees:
        society = str(employee.society or "").strip()
        key = _society_code_key(society)
        if key not in grouped:
            grouped[key] = (society, [])
        grouped[key][1].append(employee)

    desired: dict[int, str] = {}
    used: set[str] = set()
    for key in sorted(grouped, key=_employee_code_society_priority):
        society, rows = grouped[key]
        ordered = sorted(rows, key=_employee_code_sort_key)
        for employee, code in zip(
            ordered,
            _employee_code_sequence(employee_code_prefixes_for_society(society), len(ordered), used),
        ):
            if code in used:
                raise HTTPException(status_code=409, detail=f"Code employé généré en double: {code}")
            used.add(code)
            desired[int(employee.id)] = code

    updates = [
        employee for employee in employees
        if str(employee.code or "").strip().upper() != desired.get(int(employee.id), "")
        or not _employee_code_extra_matches(employee.extra, desired.get(int(employee.id), ""))
    ]
    if not updates:
        return 0

    for employee in employees:
        if int(employee.id) in desired:
            employee.code = f"TMP-RESET-{employee.id}"
    db.flush()
    for employee in employees:
        code = desired.get(int(employee.id))
        if not code:
            continue
        employee.code = code
        employee.extra = _employee_code_extra(employee.extra, code)
    try:
        db.commit()
    except SQLAlchemyError as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Correction codes employés refusée : {_postgres_error_detail(exc)}") from exc
    return len(updates)


def next_employee_code(db: Session, society: Any = None) -> str:
    if db.bind.dialect.name == "postgresql":
        db.execute(text("SELECT pg_advisory_xact_lock(20260610)"))
    used = {str(row[0] or "").upper() for row in db.query(Employee.code).all()}
    for prefix in employee_code_prefixes_for_society(society):
        for number in range(1, EMPLOYEE_CODE_SERIE_LIMIT + 1):
            code = f"{prefix}{number:02d}"
            if code not in used:
                return code
    label = ", ".join(employee_code_prefixes_for_society(society))
    raise HTTPException(status_code=409, detail=f"Série de codes employé saturée ({label} 01-{EMPLOYEE_CODE_SERIE_LIMIT})")


def next_employee_code_after_conflict(db: Session, society: Any, current_code: Any) -> str:
    prefix = str(current_code or "").strip().upper()[:1]
    allowed = employee_code_prefixes_for_society(society)
    if prefix not in allowed:
        return next_employee_code(db, society)
    used = {str(row[0] or "").upper() for row in db.query(Employee.code).all()}
    start_index = allowed.index(prefix)
    for serie_prefix in allowed[start_index:]:
        for number in range(1, EMPLOYEE_CODE_SERIE_LIMIT + 1):
            code = f"{serie_prefix}{number:02d}"
            if code not in used:
                return code
    for serie_prefix in allowed[:start_index]:
        for number in range(1, EMPLOYEE_CODE_SERIE_LIMIT + 1):
            code = f"{serie_prefix}{number:02d}"
            if code not in used:
                return code
    raise HTTPException(status_code=409, detail="Code employé indisponible")


def contract_values(employee: Employee, request: Any | None = None) -> dict[str, str]:
    extra = employee.extra if isinstance(employee.extra, dict) else {}
    request_values = getattr(request, "values", None) if request is not None else None
    request_values = request_values if isinstance(request_values, dict) else {}
    start = getattr(request, "start_date", None) if request is not None else None
    end = getattr(request, "end_date", None) if request is not None else None
    salary = getattr(request, "salary_net", None) if request is not None else None
    position = getattr(request, "position", None) if request is not None else None
    function = getattr(request, "function", None) if request is not None else None
    values: dict[str, Any] = {
        "CODE": employee.code,
        "MATRICULE": employee.code,
        "NOM": employee.last_name,
        "PRENOM": employee.first_name,
        "NOM_PRENOM": f"{employee.last_name or ''} {employee.first_name or ''}".strip(),
        "NOM_PERE": employee.father_name,
        "NOM_DU_PERE": employee.father_name,
        "NOM_MERE": employee.mother_name,
        "NOM_DE_LA_MERE": employee.mother_name,
        "ADRESSE": employee.address,
        "COMMUNE": employee.commune,
        "WILAYA": employee.wilaya,
        "NIN": employee.nin,
        "DATE_NAISSANCE": _date_str(employee.birth_date),
        "LIEU_NAISSANCE": employee.birth_place,
        "TELEPHONE": employee.phone,
        "EMAIL": employee.email,
        "POSTE": position or employee.position,
        "FONCTION": function or extra.get("fonction") or employee.position,
        "SOCIETE": employee.society,
        "TYPE_CONTRAT": getattr(request, "contract_type", None) if request is not None else employee.contract_type,
        "DATE_DEBUT": _date_str(start or employee.recruit_date),
        "DATE_FIN": _date_str(end or employee.contract_end_date),
        "DATE_CONTRAT": _date_str(date.today()),
        "DATE_RECRUTEMENT": _date_str(employee.recruit_date),
        "DATE_FIN_ESSAI": _date_str(employee.trial_end_date),
        "SALAIRE": salary if salary is not None else employee.salary_net,
        "SALAIRE_NET": salary if salary is not None else employee.salary_net,
        "CLAUSES_CONDITIONNELLES": "",
    }
    for key, value in extra.items():
        values.setdefault(str(key).upper(), value)
    for key, value in request_values.items():
        values[str(key).upper()] = value
    return {k: _date_str(v) for k, v in values.items()}


def _clause_matches(clause: ContractConditionalClause, values: dict[str, str]) -> bool:
    actual = values.get(str(clause.condition_field or "").upper(), "")
    expected = str(clause.condition_value or "")
    op = str(clause.condition_operator or "equals").lower()
    a = actual.lower().strip()
    e = expected.lower().strip()
    if op in {"equals", "=", "=="}:
        return a == e
    if op in {"contains", "contient"}:
        return e in a
    if op in {"not_equals", "!=", "different"}:
        return a != e
    return a == e


def matching_clauses(db: Session, template_id: int | None, values: dict[str, str]) -> dict[str, str]:
    stmt = select(ContractConditionalClause).where(ContractConditionalClause.active == 1)
    if template_id is not None:
        stmt = stmt.where(or_(ContractConditionalClause.template_id == template_id, ContractConditionalClause.template_id.is_(None)))
    clauses = db.execute(stmt.order_by(ContractConditionalClause.id)).scalars().all()
    grouped: dict[str, list[str]] = {}
    for clause in clauses:
        if _clause_matches(clause, values):
            key = str(clause.placeholder or "CLAUSES_CONDITIONNELLES").upper()
            grouped.setdefault(key, []).append(clause.content)
    return {key: "\n\n".join(parts) for key, parts in grouped.items()}


def _request_values(request: Any) -> dict[str, Any]:
    values = getattr(request, "values", None)
    return values if isinstance(values, dict) else {}


def _is_aps_template_request(request: Any, position: str, function: str) -> bool:
    values = _request_values(request)
    marker = str(values.get("USE_APS_TEMPLATE") or values.get("MODELE") or "").strip().upper()
    haystack = " ".join(str(v or "") for v in (position, function, values.get("FONCTION"), values.get("POSTE"))).upper()
    return marker in {"1", "APS", "CONTRAT_APS_2026"} or ("PREVENTION" in haystack and "SECURITE" in haystack)


def _fallback_aps_contract_template() -> Any:
    if not DEFAULT_APS_TEMPLATE_PATH.exists():
        return None
    return SimpleNamespace(
        id=None,
        code="CONTRAT_APS_2026",
        title="Contrat APS 2026",
        contract_type="CDD",
        position="AGENT DE PREVENTION ET DE SECURITE",
        function="AGENT DE PREVENTION ET DE SECURITE",
        file_name=DEFAULT_APS_TEMPLATE_PATH.name,
        mime_type=DOCX_MIME,
        docx_content=DEFAULT_APS_TEMPLATE_PATH.read_bytes(),
        placeholders={"items": extract_docx_placeholders(DEFAULT_APS_TEMPLATE_PATH.read_bytes())},
        active=1,
    )


def find_contract_template(db: Session, request: Any, employee: Employee) -> Any:
    if request.template_id:
        row = db.get(ContractTemplate, request.template_id)
        if not row or row.active != 1:
            raise HTTPException(status_code=404, detail="Modèle de contrat introuvable ou inactif")
        return row
    contract_type = _clean_contract_type(request.contract_type or employee.contract_type)
    position = request.position or employee.position or ""
    function = request.function or (employee.extra or {}).get("fonction") or position
    if _request_values(request).get("USE_APS_TEMPLATE"):
        fallback = _fallback_aps_contract_template()
        if fallback:
            return fallback
    stmt = select(ContractTemplate).where(ContractTemplate.active == 1, ContractTemplate.contract_type == contract_type)
    candidates = db.execute(stmt.order_by(ContractTemplate.id.desc())).scalars().all()
    if not candidates:
        if _is_aps_template_request(request, position, function):
            fallback = _fallback_aps_contract_template()
            if fallback:
                return fallback
        raise HTTPException(status_code=404, detail="Aucun modèle actif pour ce type de contrat")
    def score(t: ContractTemplate) -> int:
        points = 0
        if t.position and position and t.position.lower() == position.lower():
            points += 4
        if t.function and function and t.function.lower() == str(function).lower():
            points += 4
        if not t.position and not t.function:
            points += 1
        return points
    return sorted(candidates, key=score, reverse=True)[0]


def _replace_text_in_paragraph(paragraph, values: dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return
    replaced = PLACEHOLDER_RE.sub(lambda m: values.get(m.group(1).upper(), ""), full_text)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_text_in_table(table, values: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, values)
            for nested_table in cell.tables:
                _replace_text_in_table(nested_table, values)


def _replace_text_in_document_part(part, values: dict[str, str]) -> None:
    for paragraph in part.paragraphs:
        _replace_text_in_paragraph(paragraph, values)
    for table in part.tables:
        _replace_text_in_table(table, values)


def render_docx(template: ContractTemplate, values: dict[str, str]) -> bytes:
    try:
        from docx import Document as WordDocument
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="Dépendance python-docx manquante") from exc
    doc = WordDocument(BytesIO(template.docx_content))
    _replace_text_in_document_part(doc, values)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            _replace_text_in_document_part(part, values)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        source = tmp_path / "contrat.docx"
        source.write_bytes(docx_bytes)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(source)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=30,
            )
        except FileNotFoundError as exc:
            raise HTTPException(status_code=501, detail="Conversion PDF indisponible: LibreOffice/soffice non installé sur le serveur") from exc
        except subprocess.SubprocessError as exc:
            raise HTTPException(status_code=500, detail="Conversion PDF échouée") from exc
        pdf = tmp_path / "contrat.pdf"
        if not pdf.exists():
            raise HTTPException(status_code=500, detail="PDF non généré")
        return pdf.read_bytes()


def generate_contract(db: Session, request: Any, user: Any | None = None) -> GeneratedContract:
    employee = get_or_404(db, Employee, request.employee_id)
    template = find_contract_template(db, request, employee)
    values = contract_values(employee, request)
    values.update(matching_clauses(db, template.id, values))
    docx_bytes = render_docx(template, values)
    output_format = (request.output_format or "docx").lower()
    if output_format == "pdf":
        file_content = convert_docx_to_pdf(docx_bytes)
        mime_type = PDF_MIME
        ext = "pdf"
    else:
        output_format = "docx"
        file_content = docx_bytes
        mime_type = DOCX_MIME
        ext = "docx"
    now = datetime.utcnow()
    reference = f"CTR-{now:%Y%m%d}-{employee.code}-{int(now.timestamp())}"
    safe_name = f"{reference}-{template.contract_type}-{employee.last_name}-{employee.first_name}".replace(" ", "_")
    row = GeneratedContract(
        employee_id=employee.id,
        template_id=template.id,
        reference=reference,
        title=template.title,
        contract_type=request.contract_type or employee.contract_type or template.contract_type,
        position=request.position or employee.position,
        start_date=request.start_date or employee.recruit_date,
        end_date=request.end_date or employee.contract_end_date,
        output_format=output_format,
        file_name=f"{safe_name}.{ext}",
        mime_type=mime_type,
        file_content=file_content,
        values=values,
        generated_by=getattr(user, "username", None),
        status="genere",
    )
    db.add(row)
    contract = Contract(
        employee_id=employee.id,
        contract_type=row.contract_type,
        position=row.position,
        start_date=row.start_date,
        end_date=row.end_date,
        trial_end_date=employee.trial_end_date,
        salary_net=_parse_salary(values.get("SALAIRE_NET") or employee.salary_net),
        status="actif",
        template_code=template.code,
        content=f"Contrat généré automatiquement: {reference}",
    )
    db.add(contract)
    db.flush()
    row.contract_id = contract.id
    document = Document(
        owner_type="employee",
        owner_id=employee.id,
        label=f"Contrat généré - {row.contract_type}",
        file_name=row.file_name,
        file_path=f"generated_contract:{reference}",
        mime_type=row.mime_type,
        uploaded_by=getattr(user, "username", None),
    )
    db.add(document)
    db.commit()
    db.refresh(row)
    return row


def generate_contract_from_form(db: Session, request: Any, user: Any | None = None) -> GeneratedContract:
    first_name = str(request.first_name or "").strip()
    last_name = str(request.last_name or "").strip()
    if len(first_name) < 2 or len(last_name) < 2:
        raise HTTPException(status_code=422, detail="Nom et prénom obligatoires")
    extra = {
        "nomPere": request.father_name or "",
        "nomMere": request.mother_name or "",
        "numeroCnas": getattr(request, "numero_cnas", None) or "",
        "lieuTravail": request.work_place or "",
        "client": getattr(request, "client", None) or "",
        "motifRecrutement": request.recruitment_reason or "",
        "detailSalaire": request.salary_details or "",
        "dureeContrat": getattr(request, "contract_duration", None) or "",
    }
    employee = None
    nin_raw = str(request.nin or "").strip()
    nin = nin_raw if re.fullmatch(r"\d{10}", nin_raw) else ""
    code = str(getattr(request, "matricule", "") or "").strip()
    if nin:
        employee = db.execute(select(Employee).where(Employee.nin == nin)).scalar_one_or_none()
    if employee is None and code:
        employee = db.execute(select(Employee).where(Employee.code == code)).scalar_one_or_none()
    if employee is None:
        employee = Employee(
            code=code or next_employee_code(db, request.society),
            first_name=first_name,
            last_name=last_name,
            father_name=request.father_name,
            mother_name=request.mother_name,
            nin=nin or None,
            birth_date=request.birth_date,
            birth_place=request.birth_place,
            phone=getattr(request, "phone", None),
            address=getattr(request, "address", None),
            commune=getattr(request, "commune", None),
            wilaya=getattr(request, "wilaya", None),
            position=request.position or request.work_place,
            society=request.society,
            status="actif",
            contract_type=_clean_contract_type(request.contract_type),
            salary_net=float(request.salary_net or 0),
            recruit_date=request.start_date,
            contract_end_date=request.end_date,
            extra=extra,
        )
        db.add(employee)
        db.flush()
    else:
        employee.first_name = first_name
        employee.last_name = last_name
        if code:
            employee.code = code
        employee.father_name = request.father_name
        employee.mother_name = request.mother_name
        employee.birth_date = request.birth_date
        employee.birth_place = request.birth_place
        employee.phone = getattr(request, "phone", None)
        employee.address = getattr(request, "address", None)
        employee.commune = getattr(request, "commune", None)
        employee.wilaya = getattr(request, "wilaya", None)
        employee.position = request.position or request.work_place
        employee.society = request.society
        employee.contract_type = _clean_contract_type(request.contract_type or employee.contract_type)
        employee.salary_net = float(request.salary_net or employee.salary_net or 0)
        employee.recruit_date = request.start_date
        employee.contract_end_date = request.end_date
        employee.extra = {**(employee.extra if isinstance(employee.extra, dict) else {}), **extra}
        db.flush()
    generated_request = SimpleNamespace(
        employee_id=employee.id,
        template_id=request.template_id,
        contract_type=request.contract_type,
        position=request.position or request.work_place,
        function=request.position or request.work_place,
        start_date=request.start_date,
        end_date=request.end_date,
        salary_net=request.salary_net,
        output_format=request.output_format or "docx",
        values={
            **(request.values if isinstance(getattr(request, "values", None), dict) else {}),
            "CODE": code or employee.code or "",
            "MATRICULE": code or employee.code or "",
            "DATE_CONTRAT": date.today(),
            "CNAS": getattr(request, "numero_cnas", None) or "",
            "NUMERO_CNAS": getattr(request, "numero_cnas", None) or "",
            "TELEPHONE": getattr(request, "phone", None) or "",
            "ADRESSE": getattr(request, "address", None) or "",
            "LIEU_TRAVAIL": request.work_place or "",
            "CLIENT": getattr(request, "client", None) or "",
            "WILAYA": getattr(request, "wilaya", None) or "",
            "COMMUNE": getattr(request, "commune", None) or "",
            "DUREE_CONTRAT": (request.values.get("DUREE_CONTRAT") if isinstance(getattr(request, "values", None), dict) else None) or getattr(request, "contract_duration", None) or "",
            "MOTIF_RECRUTEMENT": request.recruitment_reason or "",
            "DETAIL_SALAIRE": request.salary_details or "",
        },
    )
    return generate_contract(db, generated_request, user)


def cleanup_base64_photos(db: Session) -> int:
    changed = 0
    for row in db.execute(select(Candidate)).scalars().all():
        if isinstance(row.data, dict):
            cleaned = normalize_photo_fields(row.data, fallback=str(row.id))
            if cleaned != row.data:
                row.data = cleaned
                changed += 1
    for row in db.execute(select(Employee)).scalars().all():
        if isinstance(row.extra, dict):
            cleaned = normalize_photo_fields(row.extra, fallback=row.code or str(row.id))
            if cleaned != row.extra:
                row.extra = cleaned
                changed += 1
    if changed:
        db.commit()
    return changed
