"""Couverture COMPLÈTE du module DRH (backend) — vrais endpoints, vraie base, sans mock.

Palier 1 : DRH. Employés, candidats + workflow recrutement, contrats, congés,
sanctions, documents, clauses conditionnelles, fiche de position, dashboard.
"""
from datetime import date, timedelta


# ── Helpers ───────────────────────────────────────────────────────────────────

def _emp(client, h, code, **kw):
    body = {"code": code, "first_name": kw.get("fn", f"E{code}"), "last_name": kw.get("ln", "Test"),
            "society": kw.get("society", "Iron Global Securite"), "status": kw.get("status", "actif"),
            "contract_type": kw.get("ct", "CDD")}
    r = client.post("/api/drh/employees", headers=h, json=body)
    assert r.status_code in (200, 201), r.text
    return r.json()


def _cand(client, h, first="Jamel", last="Cand", society="Iron Global Securite", **extra):
    body = {"first_name": first, "last_name": last, "society": society,
            "desired_position": "AGENT DE SECURITE", "phone": "0550000001",
            "expected_salary": 40000, "status": "nouvelle", "data": extra.get("data", {})}
    r = client.post("/api/drh/candidates", headers=h, json=body)
    assert r.status_code in (200, 201), r.text
    return r.json()["data"]


def test_candidate_contact_duplicate_warning(client, auth_headers):
    candidate = client.post(
        "/api/drh/candidates",
        headers=auth_headers,
        json={
            "first_name": "Amine",
            "last_name": "Doublon",
            "society": "Iron Global Securite",
            "desired_position": "Agent",
            "phone": "+213 770 12 34 56",
            "email": "Amine.Doublon@example.com",
            "status": "nouvelle",
            "data": {},
        },
    )
    assert candidate.status_code in (200, 201), candidate.text
    candidate_id = candidate.json()["data"]["id"]

    duplicate = client.get(
        "/api/drh/candidates/contact-duplicates",
        headers=auth_headers,
        params={"phone": "0770 12 34 56", "email": "amine.doublon@EXAMPLE.COM"},
    )
    assert duplicate.status_code == 200, duplicate.text
    payload = duplicate.json()
    assert payload["phone_exists"] is True
    assert payload["email_exists"] is True
    assert payload["duplicates"][0]["id"] == candidate_id
    assert set(payload["duplicates"][0]["fields"]) == {"telephone", "email"}

    excluded = client.get(
        "/api/drh/candidates/contact-duplicates",
        headers=auth_headers,
        params={
            "phone": "0770123456",
            "email": "amine.doublon@example.com",
            "exclude_candidate_id": candidate_id,
        },
    )
    assert excluded.status_code == 200, excluded.text
    assert excluded.json()["duplicates"] == []


def test_candidate_convocation_email_is_sent_and_traced(client, auth_headers, monkeypatch):
    candidate = _cand(
        client,
        auth_headers,
        first="Nadia",
        last="Convoquee",
        data={},
    )
    client.put(
        f"/api/drh/candidates/{candidate['id']}",
        headers=auth_headers,
        json={"email": "nadia.convocation@example.com"},
    )
    sent = {}

    def fake_send(**kwargs):
        sent.update(kwargs)

    monkeypatch.setattr("app.modules.drh.routes.send_candidate_convocation_email", fake_send)
    response = client.post(
        f"/api/drh/candidates/{candidate['id']}/convocation-email",
        headers=auth_headers,
        json={"date": "2026-09-10", "heure": "09:30", "lieu": "Siège", "motif": "Entretien"},
    )
    assert response.status_code == 200, response.text
    result = response.json()["data"]
    assert result["email_sent"] is True
    assert result["delivery"]["sender"] == "adm.conv@irongs.com"
    assert result["delivery"]["recipient"] == "nadia.convocation@example.com"
    assert sent["recipient"] == "nadia.convocation@example.com"
    assert sent["location"] == "Siège"


def test_convocation_email_copies_administration(monkeypatch):
    from app.modules.drh import convocation_email

    captured = {}

    class FakeSmtp:
        def __init__(self, *args, **kwargs):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

        def login(self, *args):
            pass

        def send_message(self, message):
            captured["message"] = message

    monkeypatch.setattr(convocation_email.settings, "smtp_host", "smtp.example.com")
    monkeypatch.setattr(convocation_email.settings, "smtp_use_ssl", True)
    monkeypatch.setattr(convocation_email.settings, "smtp_username", None)
    monkeypatch.setattr(convocation_email.settings, "smtp_password", None)
    monkeypatch.setattr(convocation_email.smtplib, "SMTP_SSL", FakeSmtp)
    convocation_email.send_candidate_convocation_email(
        recipient="candidat@example.com",
        candidate_name="Nadia Test",
        date="2026-09-10",
        time="09:30",
        location="Siège",
        purpose="Entretien",
    )
    message = captured["message"]
    assert message["To"] == "candidat@example.com"
    assert message["Bcc"] == "adm.conv@irongs.com"
    assert "adm.conv@irongs.com" in message["From"]


def test_transmitted_candidates_remain_visible_in_recruitment_list(client, auth_headers, db):
    from app.modules.drh.models import Candidate

    row = Candidate(
        first_name="Nadia", last_name="VISIBLE", phone="0666112233",
        society="Iron Global Securite", status="a_contractualiser",
        data={"statut": "a_contractualiser", "avisDecision": "Favorable"},
    )
    db.add(row)
    db.commit()
    response = client.get(
        "/api/drh/candidates/page",
        headers=auth_headers,
        params={"mode": "new", "society": "Iron Global Securite", "page": 1, "page_size": 100},
    )
    assert response.status_code == 200, response.text
    assert any(item["id"] == row.id and item["last_name"] == "VISIBLE" for item in response.json()["items"])

    row.data = {**row.data, "removedFromRecruitmentAt": "2026-09-02T21:30:00"}
    db.commit()
    hidden = client.get(
        "/api/drh/candidates/page",
        headers=auth_headers,
        params={"mode": "new", "society": "Iron Global Securite", "page": 1, "page_size": 100},
    )
    assert hidden.status_code == 200, hidden.text
    assert all(item["id"] != row.id for item in hidden.json()["items"])


# 7 sections visibles de la fiche de position (ordre imposé par le service)
_SECTIONS = ["identification", "militaire", "poste",
             "avis", "contact", "habilitations", "experience"]


def _full_candidate_data(nom="RECRUE", prenom="Karim"):
    """Données complètes couvrant tous les champs obligatoires des 7 sections."""
    identity_number = sum((index + 1) * ord(char) for index, char in enumerate(f"{nom}:{prenom}"))
    return {
        "nom": nom, "prenom": prenom,
        "dateNaissance": "1990-05-15", "lieuNaissance": "Alger",
        "sexe": "M", "nomPere": "Ahmed", "nomMere": "Fatima",
        "nin": f"{identity_number:010d}"[-10:], "situation": "celibataire", "source": "spontanee",
        "posteSouhaite": "AGENT DE SECURITE", "telephone": "0550112233",
        "avisDecision": "favorable", "avisDate": "2026-01-10",
        "avisRecruteur": "DRH", "avisCommentaire": "RAS",
        "adresse": "Rue 1", "commune": "Bab Ezzouar", "wilaya": "Alger",
        "contactUrgenceLien": "epouse", "contactUrgenceNom": "Sara",
        "contactUrgenceTel": "0550998877",
    }


def _make_reserve_candidate(client, h, nom="RECRUE", prenom="Karim"):
    """Crée un candidat puis le fait passer 'réserve' via validate-final (fiche validée)."""
    data = _full_candidate_data(nom, prenom)
    body = {"first_name": prenom, "last_name": nom, "society": "Iron Global Securite",
            "desired_position": "AGENT DE SECURITE", "phone": data["telephone"],
            "status": "nouvelle", "data": data}
    r = client.post("/api/drh/candidates", headers=h, json=body)
    assert r.status_code in (200, 201), r.text
    cid = r.json()["data"]["id"]
    for section in _SECTIONS:
        checked = client.post(
            f"/api/drh/candidates/validate-section?section={section}&candidate_id={cid}",
            headers=h,
            json=body,
        )
        assert checked.status_code == 200, checked.text
        body["data"]["sectionValidations"] = checked.json()["data"]["sectionValidations"]
    refused = client.post(f"/api/drh/candidates/{cid}/validate-final", headers=h, json={"validation_password": "incorrect-password"})
    assert refused.status_code == 401, refused.text
    fin = client.post(f"/api/drh/candidates/{cid}/validate-final", headers=h, json={"validation_password": "test-validation-password"})
    assert fin.status_code == 200, fin.text
    assert fin.json()["data"]["status"] == "reserve"
    return cid


# ── Dashboard ─────────────────────────────────────────────────────────────────

def test_drh_dashboard(client, auth_headers):
    r = client.get("/api/drh/dashboard", headers=auth_headers)
    assert r.status_code == 200
    assert isinstance(r.json(), dict)


# ── Employés (CRUD complet) ───────────────────────────────────────────────────

def test_employee_full_crud(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_E1", fn="Ali", ln="Crud")
    emp_id = emp.get("id") or emp.get("backendId")

    # GET one
    got = client.get(f"/api/drh/employees/{emp_id}", headers=auth_headers)
    assert got.status_code == 200
    assert (got.json().get("id") == emp_id)

    # UPDATE
    upd = client.put(f"/api/drh/employees/{emp_id}", headers=auth_headers, json={
        "code": "DRH_E1", "first_name": "Ali", "last_name": "Modifie",
        "society": "Iron Global Securite", "status": "actif", "phone": "0660000000",
    })
    assert upd.status_code == 200, upd.text
    assert upd.json()["last_name"].upper() == "MODIFIE"

    # DELETE
    dele = client.delete(f"/api/drh/employees/{emp_id}", headers=auth_headers)
    assert dele.status_code in (200, 204)
    assert client.get(f"/api/drh/employees/{emp_id}", headers=auth_headers).status_code == 404


def test_employee_get_404(client, auth_headers):
    assert client.get("/api/drh/employees/99999999", headers=auth_headers).status_code == 404


def test_employees_page_pagination(client, auth_headers):
    # page_size est borné à [5, 100] côté service
    for i in range(7):
        _emp(client, auth_headers, f"DRH_PG{i}")
    r = client.get("/api/drh/employees/page?page=1&page_size=5", headers=auth_headers)
    assert r.status_code == 200
    data = r.json()
    assert "items" in data and "total" in data and "pages" in data
    assert len(data["items"]) <= 5
    assert data["total"] >= 7
    assert data["pages"] >= 2


def test_employees_page_search(client, auth_headers):
    _emp(client, auth_headers, "DRH_SRCH", fn="Rechercheunique")
    r = client.get("/api/drh/employees/page?q=Rechercheunique", headers=auth_headers)
    assert r.status_code == 200
    assert any("RECHERCHEUNIQUE" in (e.get("first_name") or "").upper() for e in r.json()["items"])


def _site(client, h, name, society="Iron Global Securite"):
    r = client.post("/api/ops/sites", headers=h, json={
        "name": name, "indicatif": name[:3].upper(), "active": 1,
        "equipment_plan": {"societe": society},
    })
    assert r.status_code in (200, 201), r.text
    return r.json()["id"]


def _assign(client, h, employee_id, site_id, position=None):
    r = client.post("/api/ops/assignments", headers=h, json={
        "employee_id": int(employee_id), "site_id": int(site_id),
        "group_code": "A", "position": position, "start_date": "2026-01-01", "active": 1,
    })
    assert r.status_code in (200, 201), r.text
    return r.json()


def test_employees_list_includes_current_assignment(client, auth_headers):
    """La liste /drh/employees doit inclure le site/poste courant (jointure serveur),
    sans que le client ait besoin de reconstruire ce lien lui-même côté navigateur."""
    emp = _emp(client, auth_headers, "DRH_AFF1")
    site_id = _site(client, auth_headers, "Site Jointure Employe")
    _assign(client, auth_headers, emp["id"], site_id, position="Agent de sécurité")

    rows = client.get("/api/drh/employees", headers=auth_headers).json()
    row = next(r for r in rows if r["id"] == emp["id"])
    assert row["current_site_id"] == site_id
    assert row["current_site_name"].upper() == "SITE JOINTURE EMPLOYE"
    assert row["current_position"].upper() == "AGENT DE SÉCURITÉ"
    assert row["current_group_code"] == "A"

    paged = client.get("/api/drh/employees/page?page_size=100", headers=auth_headers).json()
    prow = next(r for r in paged["items"] if r["id"] == emp["id"])
    assert prow["current_site_name"].upper() == "SITE JOINTURE EMPLOYE"


def test_employees_list_current_assignment_switches_to_new_site(client, auth_headers):
    """Après réaffectation vers un nouveau site, le site précédent (désactivé) ne doit
    plus jamais réapparaître dans le site courant — c'est exactement le bug (site
    "fantôme" qui revient) qu'on a chassé côté client avant ce changement serveur."""
    emp = _emp(client, auth_headers, "DRH_AFF2")
    old_site = _site(client, auth_headers, "Ancien Site")
    new_site = _site(client, auth_headers, "Nouveau Site")
    _assign(client, auth_headers, emp["id"], old_site)
    _assign(client, auth_headers, emp["id"], new_site)

    rows = client.get("/api/drh/employees", headers=auth_headers).json()
    row = next(r for r in rows if r["id"] == emp["id"])
    assert row["current_site_id"] == new_site
    assert row["current_site_name"].upper() == "NOUVEAU SITE"


def test_employees_list_no_assignment_has_null_current_site(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_NOAFF")
    rows = client.get("/api/drh/employees", headers=auth_headers).json()
    row = next(r for r in rows if r["id"] == emp["id"])
    assert row["current_site_id"] is None
    assert row["current_site_name"] is None


def test_employee_fiche_position(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_FICHE")
    emp_id = emp.get("id") or emp.get("backendId")
    r = client.get(f"/api/drh/employees/{emp_id}/fiche-position", headers=auth_headers)
    assert r.status_code == 200
    assert isinstance(r.json(), dict)


def test_repair_employee_codes_requires_system_admin(client, auth_headers):
    """repair-codes est réservé à un admin système (token admin_system) — sinon 403."""
    r = client.post("/api/drh/employees/repair-codes", headers=auth_headers)
    assert r.status_code == 403, r.text


def test_flatten_employee_extra_requires_system_admin(client, auth_headers):
    """flatten-extra est réservé à un admin système (token admin_system) — sinon 403."""
    r = client.post("/api/drh/employees/flatten-extra", headers=auth_headers)
    assert r.status_code == 403, r.text


def test_rename_poste_agent_securite_requires_system_admin(client, auth_headers):
    """postes/rename-agent-securite est réservé à un admin système — sinon 403."""
    r = client.post("/api/drh/postes/rename-agent-securite", headers=auth_headers)
    assert r.status_code == 403, r.text


def test_rename_poste_agent_securite_renames_everywhere_and_is_idempotent(db):
    """Renomme le poste libre "Agent de sécurité" (toute casse/accent) vers le libellé
    officiel du catalogue, sur Employee.position, les champs legacy imbriqués de
    Employee.extra, et Candidate.desired_position. Un 2e passage ne change plus rien."""
    from app.modules.irongs.sql_bridge import rename_poste_agent_securite
    from app.modules.irongs.models import Position
    from app.modules.drh.models import Employee, Candidate

    db.add(Position(name="AGENT DE PRÉVENTION ET DE SÉCURITÉ (APS)", society=None))
    emp = Employee(code="RENPOS1", first_name="Test", last_name="Rename",
                    society="Iron Global Securite", status="actif", contract_type="CDD",
                    position="Agent de securite",
                    extra={"fonction": "agent de Sécurité", "_legacy": {
                        "affectationCourante": {"poste": "AGENT DE SECURITE", "siteName": "Site X"},
                    }})
    cand = Candidate(first_name="Cand", last_name="Test", desired_position="agent de sécurité",
                      society="Iron Global Securite")
    db.add(emp)
    db.add(cand)
    db.commit()

    result = rename_poste_agent_securite(db)
    assert result["total"] > 0
    canonical = result["canonical_label"]
    assert canonical == "AGENT DE PRÉVENTION ET DE SÉCURITÉ (APS)"

    db.refresh(emp)
    db.refresh(cand)
    assert emp.position == canonical
    assert emp.extra["fonction"] == canonical
    assert emp.extra["_legacy"]["affectationCourante"]["poste"] == canonical
    assert emp.extra["_legacy"]["affectationCourante"]["siteName"] == "Site X"
    assert cand.desired_position == canonical

    second = rename_poste_agent_securite(db)
    assert second["total"] == 0


# ── Candidats (CRUD + workflow recrutement) ───────────────────────────────────

def test_candidate_crud(client, auth_headers):
    c = _cand(client, auth_headers, first="Nadia")
    cid = c.get("id")
    assert cid

    lst = client.get("/api/drh/candidates", headers=auth_headers)
    assert lst.status_code == 200 and any(x.get("id") == cid for x in lst.json())

    page = client.get("/api/drh/candidates/page?page=1&page_size=10", headers=auth_headers)
    assert page.status_code == 200 and "items" in page.json()

    upd = client.put(f"/api/drh/candidates/{cid}", headers=auth_headers, json={
        "first_name": "Nadia", "last_name": "Modifiee", "society": "Iron Global Securite",
        "desired_position": "AGENT", "status": "nouvelle",
    })
    assert upd.status_code == 200, upd.text


def test_candidate_full_recruitment_workflow(client, auth_headers):
    """Workflow COMPLET : création -> fiche validée (réserve) -> à contractualiser
    -> recrutement (embauche) -> employé + contrat créés."""
    cid = _make_reserve_candidate(client, auth_headers, nom="WORKFLOW", prenom="Karim")

    # Étape réserve -> à contractualiser
    mark = client.post(f"/api/drh/candidates/{cid}/marquer-contractualisation", headers=auth_headers)
    assert mark.status_code == 200, mark.text
    assert mark.json()["data"]["status"] == "a_contractualiser"

    # Étape recrutement : le candidat devient un employé actif + un contrat
    r = client.post(f"/api/drh/candidates/{cid}/recruit", headers=auth_headers)
    assert r.status_code == 200, r.text
    emp = r.json()["data"]
    emp_id = emp.get("id") or emp.get("backendId")
    assert emp_id
    assert (emp.get("status") or "").lower() == "actif"

    contracts = client.get(f"/api/drh/contracts?employee_id={emp_id}", headers=auth_headers)
    assert contracts.status_code == 200
    assert len(contracts.json()) >= 1, "Le recrutement doit générer un contrat"

    # Un second clic/retry réseau est idempotent : même salarié, aucun contrat en double.
    retry = client.post(f"/api/drh/candidates/{cid}/recruit", headers=auth_headers)
    assert retry.status_code == 200, retry.text
    assert retry.json()["data"]["id"] == emp_id
    contracts_after_retry = client.get(f"/api/drh/contracts?employee_id={emp_id}", headers=auth_headers)
    assert len(contracts_after_retry.json()) == len(contracts.json())

    candidate = next(row for row in client.get("/api/drh/candidates", headers=auth_headers).json() if row["id"] == cid)
    assert candidate["data"]["convertedEmployeeId"] == emp_id
    assert emp["father_name"].upper() == "AHMED"
    assert emp["birth_date"] == "1990-05-15"


def test_recruit_rejects_non_validated_candidate(client, auth_headers):
    """Un candidat dont la fiche n'est pas validée ne peut PAS être recruté (422)."""
    c = _cand(client, auth_headers, first="NonValide")
    r = client.post(f"/api/drh/candidates/{c['id']}/recruit", headers=auth_headers)
    assert r.status_code == 422, r.text


def test_recruit_requires_contractualisation_state(client, auth_headers):
    cid = _make_reserve_candidate(client, auth_headers, nom="ETAT", prenom="Valide")
    r = client.post(f"/api/drh/candidates/{cid}/recruit", headers=auth_headers)
    assert r.status_code == 422, r.text


def test_validate_section_persists_on_existing_candidate(client, auth_headers):
    data = _full_candidate_data(nom="PERSISTE", prenom="Section")
    data.pop("sectionValidations", None)
    c = _cand(client, auth_headers, first="Section", last="PERSISTE", data=data)
    body = {
        "first_name": "Section", "last_name": "PERSISTE",
        "society": "Iron Global Securite", "data": data,
    }
    r = client.post(
        f"/api/drh/candidates/validate-section?section=identification&candidate_id={c['id']}",
        headers=auth_headers, json=body,
    )
    assert r.status_code == 200, r.text
    saved = next(row for row in client.get("/api/drh/candidates", headers=auth_headers).json() if row["id"] == c["id"])
    assert saved["data"]["sectionValidations"]["identification"]["by"]


def test_candidate_accepts_18_digit_nin(client, auth_headers):
    data = _full_candidate_data(nom="NINLONG", prenom="Dixhuit")
    data["nin"] = "123456789012345678"
    r = client.post("/api/drh/candidates", headers=auth_headers, json={
        "first_name": "Dixhuit", "last_name": "NINLONG",
        "society": "Iron Global Securite", "status": "nouvelle", "data": data,
    })
    assert r.status_code in (200, 201), r.text


def test_candidate_rejects_forged_section_validations(client, auth_headers):
    data = _full_candidate_data(nom="FORGE", prenom="Tampon")
    data["sectionValidations"] = {section: {"by": "browser", "at": "2026-08-11T00:00:00"} for section in _SECTIONS}
    data["fichePositionValidee"] = True
    candidate = _cand(client, auth_headers, first="Tampon", last="FORGE", data=data)
    result = client.post(f"/api/drh/candidates/{candidate['id']}/validate-final", headers=auth_headers, json={"validation_password": "test-validation-password"})
    assert result.status_code == 422, result.text
    assert "sections non validées" in result.text


def test_candidate_duplicate_nin_is_rejected_early(client, auth_headers):
    first_data = _full_candidate_data(nom="DOUBLON", prenom="Premier")
    _cand(client, auth_headers, first="Premier", last="DOUBLON", data=first_data)
    second_data = _full_candidate_data(nom="AUTRE", prenom="Second")
    second_data["nin"] = first_data["nin"]
    response = client.post("/api/drh/candidates", headers=auth_headers, json={
        "first_name": "Second", "last_name": "AUTRE", "society": "Iron Global Securite",
        "status": "nouvelle", "data": second_data,
    })
    assert response.status_code == 409, response.text
    assert "même NIN" in response.text


def test_candidate_api_rejects_non_recruitment_role(client, restricted_headers):
    response = client.get("/api/drh/candidates/page", headers=restricted_headers)
    assert response.status_code == 403, response.text


def test_marquer_contractualisation_accepts_non_reserve_candidate(client, auth_headers):
    """marquer-contractualisation accepte un candidat qui n'est pas en réserve : c'est le
    chemin normal des candidats transmis depuis recrute.html, qui ne passent jamais par le
    workflow interne "réserve" d'ATLAS avant d'être envoyés en contractualisation, à condition
    que la décision du recruteur soit favorable."""
    c = _cand(client, auth_headers, first="PasEnReserve", data={"avisDecision": "Favorable"})
    r = client.post(f"/api/drh/candidates/{c['id']}/marquer-contractualisation", headers=auth_headers)
    assert r.status_code == 200, r.text
    assert r.json()["data"]["status"] == "a_contractualiser"


def test_marquer_contractualisation_rejects_non_favorable_decisions(client, auth_headers):
    for index, decision in enumerate(("Instance", "Défavorable")):
        candidate = _cand(client, auth_headers, first=f"Decision{index}", data={"avisDecision": decision})
        response = client.post(
            f"/api/drh/candidates/{candidate['id']}/marquer-contractualisation",
            headers=auth_headers,
        )
        assert response.status_code == 422, response.text
        assert "Favorable" in response.text


def test_marquer_contractualisation_rejects_archived_or_embauche(client, auth_headers):
    """marquer-contractualisation refuse un dossier déjà archivé ou déjà transformé en
    employé (409) — c'est la seule restriction de statut réellement appliquée."""
    c = _cand(client, auth_headers, first="DejaArchive")
    patch = client.put(f"/api/drh/candidates/{c['id']}", headers=auth_headers, json={"status": "archive"})
    assert patch.status_code == 200, patch.text
    r = client.post(f"/api/drh/candidates/{c['id']}/marquer-contractualisation", headers=auth_headers)
    assert r.status_code == 409, r.text


def test_validate_section_enforces_order(client, auth_headers):
    """validate-section impose l'ordre : valider 'poste' avant les précédentes est refusé."""
    data = _full_candidate_data(nom="ORDRE", prenom="Test")
    data.pop("sectionValidations", None)  # aucune section validée
    body = {"first_name": "Test", "last_name": "ORDRE", "society": "Iron Global Securite",
            "desired_position": "AGENT", "status": "nouvelle", "data": data}
    r = client.post("/api/drh/candidates/validate-section?section=poste", headers=auth_headers, json=body)
    assert r.status_code == 422, r.text
    assert "précédente" in r.text or "precedente" in r.text.lower()


def test_candidate_delete(client, auth_headers):
    c = _cand(client, auth_headers, first="ASupprimer")
    cid = c.get("id")
    r = client.delete(f"/api/drh/candidates/{cid}", headers=auth_headers)
    assert r.status_code == 200, r.text


def test_candidate_validate_section(client, auth_headers):
    r = client.post("/api/drh/candidates/validate-section?section=identite", headers=auth_headers, json={
        "first_name": "Valide", "last_name": "Section", "society": "Iron Global Securite",
        "desired_position": "AGENT", "status": "nouvelle", "data": {},
    })
    # L'endpoint répond (succès ou erreurs de validation structurées)
    assert r.status_code in (200, 400, 422), r.text


# ── Contrats ──────────────────────────────────────────────────────────────────

def test_contract_crud(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_CTR")
    emp_id = emp.get("id") or emp.get("backendId")
    c = client.post("/api/drh/contracts", headers=auth_headers, json={
        "employee_id": emp_id, "contract_type": "CDD", "start_date": "2026-01-01",
        "end_date": "2026-12-31", "salary_net": 45000, "status": "actif",
    })
    assert c.status_code in (200, 201), c.text
    contract_id = c.json()["id"]

    upd = client.put(f"/api/drh/contracts/{contract_id}", headers=auth_headers, json={"salary_net": 50000})
    assert upd.status_code == 200, upd.text
    assert upd.json()["salary_net"] == 50000

    lst = client.get(f"/api/drh/contracts?employee_id={emp_id}", headers=auth_headers)
    assert lst.status_code == 200 and len(lst.json()) >= 1


# ── Congés (workflow approbation) ─────────────────────────────────────────────

def test_leave_workflow_approve(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_LV1")
    emp_id = emp.get("id") or emp.get("backendId")
    lv = client.post("/api/drh/leaves", headers=auth_headers, json={
        "employee_id": emp_id, "leave_type": "conge",
        "start_date": str(date.today()), "end_date": str(date.today() + timedelta(days=5)),
        "reason": "Congé annuel",
    })
    assert lv.status_code in (200, 201), lv.text
    leave_id = lv.json()["id"]

    appr = client.post(f"/api/drh/leaves/{leave_id}/approve", headers=auth_headers)
    assert appr.status_code == 200
    assert appr.json()["status"] == "approuve"


def test_leave_workflow_refuse(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_LV2")
    emp_id = emp.get("id") or emp.get("backendId")
    lv = client.post("/api/drh/leaves", headers=auth_headers, json={
        "employee_id": emp_id, "start_date": str(date.today()),
        "end_date": str(date.today() + timedelta(days=2)),
    }).json()
    ref = client.post(f"/api/drh/leaves/{lv['id']}/refuse", headers=auth_headers)
    assert ref.status_code == 200 and ref.json()["status"] == "refuse"


def test_leaves_list(client, auth_headers):
    r = client.get("/api/drh/leaves", headers=auth_headers)
    assert r.status_code == 200 and isinstance(r.json(), list)


# ── Sanctions ─────────────────────────────────────────────────────────────────

def test_sanction_create_and_list(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_SANC")
    emp_id = emp.get("id") or emp.get("backendId")
    s = client.post("/api/drh/sanctions", headers=auth_headers, json={
        "employee_id": emp_id, "infraction_date": str(date.today()),
        "fault": "Retard répété", "sanction_type": "avertissement", "suspension_days": 0,
    })
    assert s.status_code in (200, 201), s.text
    lst = client.get(f"/api/drh/sanctions?employee_id={emp_id}", headers=auth_headers)
    assert lst.status_code == 200 and len(lst.json()) >= 1


# ── Documents ─────────────────────────────────────────────────────────────────

def test_document_create_and_list(client, auth_headers):
    emp = _emp(client, auth_headers, "DRH_DOC")
    emp_id = emp.get("id") or emp.get("backendId")
    d = client.post("/api/drh/documents", headers=auth_headers, json={
        "owner_type": "employee", "owner_id": emp_id, "label": "Contrat signé",
        "file_name": "contrat.pdf", "mime_type": "application/pdf",
    })
    assert d.status_code in (200, 201), d.text
    lst = client.get(f"/api/drh/documents?owner_type=employee&owner_id={emp_id}", headers=auth_headers)
    assert lst.status_code == 200 and len(lst.json()) >= 1


# ── Clauses conditionnelles de contrat ────────────────────────────────────────

def test_contract_clause_crud(client, auth_headers):
    c = client.post("/api/drh/contract-clauses", headers=auth_headers, json={
        "title": "Clause zone", "condition_field": "function", "condition_operator": "equals",
        "condition_value": "AGENT DE SECURITE", "content": "L'agent effectue des rondes.", "active": 1,
    })
    assert c.status_code in (200, 201), c.text
    clause_id = c.json()["id"]

    lst = client.get("/api/drh/contract-clauses", headers=auth_headers)
    assert lst.status_code == 200 and any(x["id"] == clause_id for x in lst.json())

    upd = client.put(f"/api/drh/contract-clauses/{clause_id}", headers=auth_headers, json={
        "title": "Clause zone modifiée", "condition_field": "function", "condition_operator": "equals",
        "condition_value": "AGENT DE SECURITE", "content": "Rondes toutes les heures.", "active": 1,
    })
    assert upd.status_code == 200, upd.text

    dele = client.delete(f"/api/drh/contract-clauses/{clause_id}", headers=auth_headers)
    assert dele.status_code in (200, 204)


def test_contract_templates_list(client, auth_headers):
    r = client.get("/api/drh/contract-templates", headers=auth_headers)
    assert r.status_code == 200 and isinstance(r.json(), list)


def _minimal_docx_bytes(text="Contrat de {{NOM}} {{PRENOM}} — poste {{POSTE}}."):
    from io import BytesIO
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_preview_contract_from_form_merges_without_side_effects(client, auth_headers):
    """L'aperçu d'un modèle Word personnalise doit fusionner les donnees du formulaire
    SANS creer d'employe ni de contrat/document persistes (contrairement a la
    generation finale) — c'est l'invariant de securite du bouton Apercu."""
    upload = client.post(
        "/api/drh/contract-templates",
        headers=auth_headers,
        data={"code": "TEST_COORD", "title": "Contrat Coordinateur", "contract_type": "CDD",
              "position": "COORDINATEUR", "active": 1},
        files={"file": ("coordinateur.docx", _minimal_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert upload.status_code == 200, upload.text
    template_id = upload.json()["id"]

    employees_before = client.get("/api/drh/employees", headers=auth_headers).json()

    payload = {
        "template_id": template_id,
        "matricule": "PREVIEWTEST",
        "contract_type": "CDD",
        "first_name": "Slimane",
        "last_name": "Rouabeh",
        "position": "COORDINATEUR",
        "society": "Iron Global Securite",
        "start_date": str(date.today()),
        "end_date": str(date.today() + timedelta(days=365)),
        "salary_net": 45000,
        "values": {"FONCTION": "COORDINATEUR", "POSTE": "COORDINATEUR"},
        "output_format": "docx",
    }
    preview = client.post("/api/drh/generated-contracts/preview-from-form", headers=auth_headers, json=payload)
    assert preview.status_code == 200, preview.text
    assert preview.headers["content-type"].startswith("application/vnd.openxmlformats")
    assert len(preview.content) > 0
    from io import BytesIO
    from docx import Document
    merged_text = "\n".join(paragraph.text for paragraph in Document(BytesIO(preview.content)).paragraphs)
    assert "Rouabeh" in merged_text
    assert "Slimane" in merged_text
    assert "COORDINATEUR" in merged_text
    assert "{{NOM}}" not in merged_text and "{{PRENOM}}" not in merged_text and "{{POSTE}}" not in merged_text

    employees_after = client.get("/api/drh/employees", headers=auth_headers).json()
    assert len(employees_after) == len(employees_before), "L'aperçu ne doit jamais créer d'employé"

    generated = client.get("/api/drh/generated-contracts", headers=auth_headers)
    if generated.status_code == 200:
        codes = [g.get("employee_id") for g in generated.json()]
        assert "PREVIEWTEST" not in [str(c) for c in codes]


def test_preview_contract_from_form_requires_template_id(client, auth_headers):
    payload = {
        "first_name": "Sans", "last_name": "Modele", "contract_type": "CDD",
        "society": "Iron Global Securite", "output_format": "docx",
    }
    r = client.post("/api/drh/generated-contracts/preview-from-form", headers=auth_headers, json=payload)
    assert r.status_code == 422


def _agent_headers(client, db):
    """Compte role='agent' (consultation simple) — ne doit jamais pouvoir modifier les
    modèles de contrat ni les clauses conditionnelles (document officiel partagé)."""
    from app.core.security import hash_password
    from app.modules.auth.models import User
    if not db.query(User).filter(User.username == "testagent_contract").first():
        db.add(User(username="testagent_contract", email=None, full_name="Agent Contrat",
                     role="agent", access_level="H1", authorized_societies=[], authorized_structures=[],
                     password_hash=hash_password("testpass123"), is_active=True))
        db.commit()
    tok = client.post("/api/auth/login", json={"username": "testagent_contract", "password": "testpass123"})
    assert tok.status_code == 200, tok.text
    return {"Authorization": f"Bearer {tok.json()['access_token']}"}


def test_agent_role_cannot_write_contract_clauses(client, auth_headers, db):
    h = _agent_headers(client, db)
    r = client.post("/api/drh/contract-clauses", headers=h, json={
        "title": "Clause interdite", "condition_field": "function", "condition_operator": "equals",
        "condition_value": "X", "content": "Y", "active": 1,
    })
    assert r.status_code == 403, r.text

    c = client.post("/api/drh/contract-clauses", headers=auth_headers, json={
        "title": "Clause admin", "condition_field": "function", "condition_operator": "equals",
        "condition_value": "X", "content": "Y", "active": 1,
    })
    clause_id = c.json()["id"]
    assert client.put(f"/api/drh/contract-clauses/{clause_id}", headers=h, json={
        "title": "Modif interdite", "condition_field": "function", "condition_operator": "equals",
        "condition_value": "X", "content": "Z", "active": 1,
    }).status_code == 403
    assert client.delete(f"/api/drh/contract-clauses/{clause_id}", headers=h).status_code == 403


def test_generated_contracts_list(client, auth_headers):
    r = client.get("/api/drh/generated-contracts", headers=auth_headers)
    assert r.status_code == 200 and isinstance(r.json(), list)


# ── Logique métier du service (appels directs, vraie base) ────────────────────

def test_employee_code_prefixes_by_society():
    from app.modules.drh.service import employee_code_prefixes_for_society as prefixes
    assert prefixes("Iron Global Securite") == ["A", "B", "C"]
    assert prefixes("IRON GLOBAL SÉCURITÉ") == ["A", "B", "C"]  # accents/casse normalisés
    assert prefixes("Iron Global Solution") == ["K", "W"]
    assert prefixes("Sword Corporation") == ["S"]
    assert prefixes("Sword Construction") == ["T"]
    assert len(prefixes("Societe Inconnue")) == 26  # repli : tout l'alphabet


def test_employee_code_sequence_skips_used_codes():
    from app.modules.drh.service import _employee_code_sequence
    codes = _employee_code_sequence(["A"], 3, used={"A01", "A03"})
    assert codes == ["A02", "A04", "A05"]


def test_employee_code_sequence_saturation_raises_409():
    from fastapi import HTTPException
    import pytest
    from app.modules.drh.service import _employee_code_sequence, EMPLOYEE_CODE_SERIE_LIMIT
    full = {f"A{n:02d}" for n in range(1, EMPLOYEE_CODE_SERIE_LIMIT + 1)}
    with pytest.raises(HTTPException) as exc:
        _employee_code_sequence(["A"], 1, used=full)
    assert exc.value.status_code == 409


def test_employee_code_extra_sync():
    from app.modules.drh.service import _employee_code_extra, _employee_code_extra_matches
    extra = {"matricule": "Z99", "code": "Z99", "_legacy": {"matricule": "Z99", "code": "Z99"}}
    assert _employee_code_extra_matches(extra, "Z99") is True
    assert _employee_code_extra_matches(extra, "A01") is False
    fixed = _employee_code_extra(extra, "A01")
    assert fixed["matricule"] == "A01" and fixed["code"] == "A01"
    assert fixed["_legacy"]["matricule"] == "A01" and fixed["_legacy"]["code"] == "A01"
    assert _employee_code_extra_matches(fixed, "A01") is True


def test_next_employee_code_respects_society_prefix(db):
    from app.modules.drh.service import next_employee_code, employee_code_prefixes_for_society
    code = next_employee_code(db, "Sword Corporation")
    assert code[0] in employee_code_prefixes_for_society("Sword Corporation")
    # Le code proposé n'est jamais déjà pris
    from app.modules.drh.models import Employee
    used = {str(r[0] or "").upper() for r in db.query(Employee.code).all()}
    assert code not in used


def test_next_employee_code_after_conflict_keeps_serie(db):
    from app.modules.drh.service import next_employee_code_after_conflict
    # Un code hors série retombe sur la série de la société
    code = next_employee_code_after_conflict(db, "Sword Corporation", "Z42")
    assert code.startswith("S")
    # Un code dans la série reste dans la série
    code2 = next_employee_code_after_conflict(db, "Sword Corporation", "S01")
    assert code2.startswith("S")


def test_clause_matches_all_operators():
    from app.modules.drh.models import ContractConditionalClause
    from app.modules.drh.service import _clause_matches
    values = {"FUNCTION": "Agent de Securite"}

    def clause(op, val, field="function"):
        return ContractConditionalClause(condition_field=field, condition_operator=op,
                                         condition_value=val, placeholder="P", content="C", active=1)

    assert _clause_matches(clause("equals", "agent de securite"), values) is True   # insensible à la casse
    assert _clause_matches(clause("equals", "chauffeur"), values) is False
    assert _clause_matches(clause("contains", "securite"), values) is True
    assert _clause_matches(clause("contient", "chauffeur"), values) is False
    assert _clause_matches(clause("not_equals", "chauffeur"), values) is True
    assert _clause_matches(clause("!=", "agent de securite"), values) is False
    assert _clause_matches(clause("operateur_inconnu", "agent de securite"), values) is True  # repli = equals
    assert _clause_matches(clause("equals", "x", field="champ_absent"), values) is False


def test_contract_values_maps_employee_fields():
    from datetime import date as _date
    from app.modules.drh.models import Employee
    from app.modules.drh.service import contract_values
    emp = Employee(code="A01", first_name="Karim", last_name="Benali", society="Iron Global Securite",
                   father_name="Ahmed", mother_name="Fatima", position="AGENT",
                   birth_date=_date(1990, 5, 15), nin="1234567890", salary_net=45000,
                   extra={"fonction": "AGENT DE SECURITE"})
    v = contract_values(emp)
    assert v["CODE"] == "A01" and v["MATRICULE"] == "A01"
    assert v["NOM"] == "Benali" and v["PRENOM"] == "Karim"
    assert v["NOM_PRENOM"] == "Benali Karim"
    assert v["NOM_PERE"] == "Ahmed" and v["NOM_DE_LA_MERE"] == "Fatima"
    assert v["DATE_NAISSANCE"] == "15/05/1990"  # format contrat : jj/mm/aaaa
    assert v["FONCTION"] == "AGENT DE SECURITE"  # extra.fonction prioritaire sur position
    assert v["SALAIRE_NET"] == "45000"
    assert v["CLAUSES_CONDITIONNELLES"] == ""


def test_contract_values_request_overrides_employee():
    from app.modules.drh.models import Employee
    from app.modules.drh.service import contract_values

    class Req:
        position = "CHEF DE POSTE"
        function = "SUPERVISEUR"
        salary_net = 60000
        start_date = "2026-03-01"
        end_date = "2026-12-31"
        contract_type = "CDD"
        values = {"lieu": "Alger"}

    emp = Employee(code="B02", first_name="Sara", last_name="Amrani",
                   society="Iron Global Securite", position="AGENT", salary_net=40000)
    v = contract_values(emp, Req())
    assert v["POSTE"] == "CHEF DE POSTE"
    assert v["FONCTION"] == "SUPERVISEUR"
    assert v["SALAIRE"] == "60000"
    assert v["DATE_DEBUT"] == "2026-03-01" and v["DATE_FIN"] == "2026-12-31"
    assert v["TYPE_CONTRAT"] == "CDD"
    assert v["LIEU"] == "Alger"  # les values de la requête sont ajoutées en MAJUSCULES


def test_matching_clauses_groups_by_placeholder(client, auth_headers, db):
    from app.modules.drh.service import matching_clauses
    for i, content in enumerate(["Clause A", "Clause B"]):
        r = client.post("/api/drh/contract-clauses", headers=auth_headers, json={
            "title": f"MC{i}", "condition_field": "function", "condition_operator": "equals",
            "condition_value": "MATCHFN", "placeholder": "BLOC_TEST", "content": content, "active": 1,
        })
        assert r.status_code in (200, 201), r.text
    result = matching_clauses(db, None, {"FUNCTION": "MATCHFN"})
    assert "BLOC_TEST" in result
    assert "Clause A" in result["BLOC_TEST"] and "Clause B" in result["BLOC_TEST"]
    # Une fonction qui ne matche pas ne ramène rien pour ce bloc
    assert "BLOC_TEST" not in matching_clauses(db, None, {"FUNCTION": "AUTRE"})


def test_repair_employee_codes_renumbers_alphabetically(client, auth_headers, db):
    """repair_employee_codes_if_needed renumérote par société, ordre alphabétique,
    sans doublon, et est idempotent (2e appel = 0 correction)."""
    from app.modules.drh.service import repair_employee_codes_if_needed
    from app.modules.drh.models import Employee

    changed = repair_employee_codes_if_needed(db)
    assert isinstance(changed, int)

    rows = db.query(Employee).all()
    codes = [str(e.code or "") for e in rows]
    assert len(codes) == len(set(codes)), "Codes employés en double après réparation"
    # Chaque code est dans la série de sa société, et extra est synchronisé
    from app.modules.drh.service import employee_code_prefixes_for_society, _employee_code_extra_matches
    for e in rows:
        assert e.code[0] in employee_code_prefixes_for_society(e.society), f"{e.code} hors série pour {e.society}"
        assert _employee_code_extra_matches(e.extra, e.code), f"extra désynchronisé pour {e.code}"

    # Idempotence
    assert repair_employee_codes_if_needed(db) == 0
